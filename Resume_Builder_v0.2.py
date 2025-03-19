from docx import Document
from docx.shared import Pt, Inches
from db_manager import (
    get_employment,
    get_publications,
    get_education,
    get_person_info,
    get_skills,
    get_professional_development,
    # get_certifications,
    get_projects,
)
from job_relavancy_scorer import score_and_rank_relevance
from job_posting_scraper_ai import get_scraped_job_data

def fetch_resume_data(db_path, person_id, search_filters=None):
    """Fetches all resume-related data from the database."""
    if search_filters is None:
        search_filters = {}
    (full_name, email, linkedin, github) = get_person_info(db_path, person_id)

    return {
        "full_name": full_name,
        "email": email,
        "linkedin": linkedin,
        "github": github,
        "education": get_education(db_path, person_id),
        "employment": get_employment(
            db_path, person_id
        ),  # , search_filters["employment_filters"]),
        "publications": get_publications(db_path, person_id),
        "projects": get_projects(
            db_path, person_id, fields=["Data Science"], exclude_fields=["Personal"]
        ),
        "personal_projects": get_projects(db_path, person_id, types=["Personal"]),
        "professional_development": get_professional_development(db_path, person_id),
        "skills": get_skills(db_path, person_id),
    }


def build_resume(
    person_id,
    db_path,
    job_posting_url,
    output_file="formatted_document.docx",
    duty_filters=None,
    ):

    # Create a new document
    doc = Document()
    data = fetch_resume_data(db_path, person_id, duty_filters)
    job_data = get_scraped_job_data(db_path, job_posting_url)

    # Set document margins to 0.5 inches
    sections = doc.sections
    margin = 0.5
    for section in sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    # Add the first line with 16pt font and no paragraph spacing
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run(f"{data['full_name']}")
    run1.font.size = Pt(16)
    run1.bold = True

    # Add the second line with custom tab stops and no paragraph spacing
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    # Define tab stops
    usable_width = float(sections[0].page_width.inches) - (2.0 * margin)
    tab_stops = p2.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(0))      # Left-aligned word
    tab_stops.add_tab_stop(Inches(usable_width / 2), 1)   # Center-aligned word
    tab_stops.add_tab_stop(Inches(usable_width), 2)   # Right-aligned word

    # Add the text with tab characters
    run2 = p2.add_run(f"{data['email']}\t{data['linkedin']}\t{data['github']}\t")
    run2.font.size = Pt(10)

    # Add third and fourth line for personal statement
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    run3 = p3.add_run("Personal Statement")
    run3.font.size = Pt(12)
    run3.bold = True

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)
    tab_stops = p4.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(margin /2))  # Left-aligned word
    run4 = p4.add_run("\tPersonal Statement goes here...")
    run4.font.size = Pt(10)

    # Check on experience and ad data if necessary
    for item in data:
        if item not in ["full_name", "email", "linkedin", "github"]:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if item == "education":
                if data[item]:
                    run = para.add_run("Education")
                    run.font.size = Pt(12)
                    run.bold = True
                    tab_stops = para.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(0))  # Left-aligned word
                    tab_stops.add_tab_stop(Inches(margin / 2))  # Left-aligned word
                    tab_stops.add_tab_stop(Inches(margin))  # Left-aligned word
                    tab_stops.add_tab_stop(Inches(usable_width - margin), 2)  # Right-aligned word

                    for degree, institution, graduation_year, gpa in data["education"]:
                        run = para.add_run(f"\n\t{institution}, \t{graduation_year}\n")
                        run.font.size = Pt(10)
                        run.bold = True
                        degree_run = para.add_run(f"\t\t{degree}\t GPA: {gpa}/4.0")
                        degree_run.font.size = Pt(10)
            if item == "employment":
                if data[item]:
                    prev_company= ""
                    for position in data["employment"]:
                        (
                            company,
                            location,
                            title,
                            start_date,
                            end_date,
                            field,
                            responsibilities,
                            # fields,
                        ) = position

                        if responsibilities:
                            responsibilities = score_and_rank_relevance(job_data, position)
                        if responsibilities is None or responsibilities is []:
                            pass
                        else:
                            run = para.add_run("Employment")
                            run.font.size = Pt(12)
                            run.bold = True
                            tab_stops = para.paragraph_format.tab_stops
                            tab_stops.add_tab_stop(Inches(0))  # Left-aligned word
                            tab_stops.add_tab_stop(Inches(margin / 2))  # Left-aligned word
                            tab_stops.add_tab_stop(Inches(margin))  # Left-aligned word
                            tab_stops.add_tab_stop(Inches(usable_width - margin), 2)  # Right-aligned word

                            # Reset paragraph indentation for each new company
                            emp_para = doc.add_paragraph()
                            emp_para.paragraph_format.space_before = Pt(0)
                            emp_para.paragraph_format.space_after = Pt(0)
                            tab_stops = emp_para.paragraph_format.tab_stops
                            tab_stops.add_tab_stop(Inches(0))  # Left-aligned word
                            tab_stops.add_tab_stop(Inches(margin / 2))  # Left-aligned word
                            tab_stops.add_tab_stop(Inches(margin))  # Left-aligned word
                            tab_stops.add_tab_stop(Inches(usable_width - margin), 2)  # Right-aligned word

                            if prev_company != company:
                                # Add company with bold formatting
                                company_run = emp_para.add_run(f"\t{company}")
                                company_run.bold = True
                                # Add location if available
                                location_run = emp_para.add_run(f", {location}\n")
                                # Apply font formatting
                                company_run.font.size = Pt(10)
                                location_run.font.size = Pt(10)

                            # Add job title and dates with tab spacing
                            date_range = (
                                f"{start_date} - {end_date}"
                                if end_date
                                else f"{start_date} - Present"
                            )

                            title_run = emp_para.add_run(f"\t\t{title}\t{date_range}")
                            title_run.font.size = Pt(10)

                            # Process responsibilities as bullet points
                            if responsibilities:
                                for responsibility in responsibilities.split(';'):

                                    if responsibility.strip():
                                        # Create a properly formatted bullet point with the dash and spaces
                                        bullet_para = doc.add_paragraph(style='List Bullet')
                                        bullet_para.paragraph_format.space_before = Pt(0)
                                        bullet_para.paragraph_format.space_after = Pt(0)
                                        bullet_para.paragraph_format.left_indent = Inches(margin * 1.5)
                                        bullet_run = bullet_para.add_run(
                                            f"{responsibility.strip()}"
                                        )
                                        bullet_run.font.size = Pt(10)
                            # Update previous company
                            prev_company = company
                            # if i < len(data["employment"]):
                            #     # Add space between job entries
                            #     current_para.add_run("\n")
            elif item == "projects" or item == "personal_projects":
                if item == "projects":
                    focus = "projects"
                    section_title = "Projects"
                else:
                    focus = "personal_projects"
                    section_title = "Personal Projects"
                run = para.add_run(section_title)
                run.font.size = Pt(12)
                run.bold = True


                for project in data[focus]:
                    (
                        project_name,
                        year,
                        technologies,
                        project_link,
                        field,
                        project_type,
                        details,
                    ) = project

                    # Reset paragraph indentation for each new company
                    proj_para = doc.add_paragraph()
                    proj_para.paragraph_format.space_before = Pt(0)
                    proj_para.paragraph_format.space_after = Pt(0)
                    tab_stops = proj_para.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(0))  # Left-aligned word
                    tab_stops.add_tab_stop(Inches(margin / 2))  # Left-aligned word
                    project_run = proj_para.add_run(f"\t{project_name}")
                    project_run.bold = True
                    project_run.font.size = Pt(10)

                    if field:
                        field_run = proj_para.add_run(f", {field}, {year}")
                        field_run.font.size = Pt(10)

                    tools_para = doc.add_paragraph()
                    tools_para.paragraph_format.space_before = Pt(0)
                    tools_para.paragraph_format.space_after = Pt(0)
                    tab_stops = tools_para.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(margin / 2))  # Left-aligned word

                    tools_run = tools_para.add_run(
                        f"\tTools & Technologies: {technologies}"
                    )
                    tools_run.font.size = Pt(10)

                    if details:
                        for detail in details.split(";"):
                            if detail.strip():
                                # Create a properly formatted bullet point with the dash and spaces
                                bullet_para = doc.add_paragraph(style='List Bullet')
                                bullet_para.paragraph_format.space_before = Pt(0)
                                bullet_para.paragraph_format.space_after = Pt(0)
                                bullet_para.paragraph_format.left_indent = Inches(margin)
                                detail_run = bullet_para.add_run(f"{detail.strip()}")
                                detail_run.font.size = Pt(10)


    # Save the document
    print(f"Saving file as {output_file}")
    doc.save(output_file)

if __name__ == "__main__":
    job_url = "https://sjobs.brassring.com/TGnewUI/Search/Home/Home?partnerid=26336&siteid=5014#jobDetails=1572907_5014"
    build_resume(
        1, db_path="resume.db", job_posting_url=job_url
    )  # , filters=filters)
