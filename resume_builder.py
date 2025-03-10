from docx import Document
from db_manager import (
    get_employment,
    get_publications,
    get_education,
    get_person_info,
    get_skills,
    get_professional_development,
    # get_certifications,
    get_projects,
)
from docx.shared import Pt


def fetch_resume_data(db_path, person_id, search_filters=None):
    """Fetches all resume-related data from the database."""
    if search_filters is None:
        search_filters = {}
    (full_name, email, linkedin, github) = get_person_info(
        db_path, person_id
    )

    return {
        "full_name": full_name,
        "email": email,
        "linkedin": linkedin,
        "github": github,
        "education": get_education(db_path, person_id),
        "employment": get_employment(db_path, person_id, search_filters["employment_filters"]),
        "publications": get_publications(db_path, person_id),
        "projects": get_projects(db_path, person_id, fields=["Data Science"], exclude_fields=["Personal"]),
        "personal_projects": get_projects(db_path, person_id, types=["Personal"]),
        "professional_development": get_professional_development(db_path, person_id),
        "skills": get_skills(db_path, person_id)
    }


def replace_text_while_keeping_formatting(paragraph, key, value):
    """Replace text in a paragraph while maintaining the original formatting."""
    if key in paragraph.text:
        # Store initial formatting
        runs_with_formats = []
        for run in paragraph.runs:
            runs_with_formats.append(
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size,
                    "color": run.font.color.rgb if run.font.color else None,
                }
            )

        # Clear the paragraph
        paragraph.clear()

        # Reconstruct the paragraph with the replacement
        for run_format in runs_with_formats:
            original_text = run_format["text"]
            new_text = original_text.replace(key, value)

            if new_text != original_text:  # Text was replaced
                new_run = paragraph.add_run(new_text)
                # Apply stored formatting
                new_run.bold = run_format["bold"]
                new_run.italic = run_format["italic"]
                new_run.underline = run_format["underline"]
                new_run.font.name = run_format["font_name"]
                if run_format["font_size"]:
                    new_run.font.size = run_format["font_size"]
                if run_format["color"]:
                    new_run.font.color.rgb = run_format["color"]
            else:  # Text wasn't replaced, keep original
                new_run = paragraph.add_run(original_text)
                # Apply stored formatting
                new_run.bold = run_format["bold"]
                new_run.italic = run_format["italic"]
                new_run.underline = run_format["underline"]
                new_run.font.name = run_format["font_name"]
                if run_format["font_size"]:
                    new_run.font.size = run_format["font_size"]
                if run_format["color"]:
                    new_run.font.color.rgb = run_format["color"]


def replace_text_with_tabs(paragraph, replacements):
    """Replace tab-separated placeholders while preserving formatting."""
    parts = paragraph.text.split("\t")  # Split by tab

    # Replace placeholders
    new_parts = [replacements.get(part, part) for part in parts]

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # Rebuild the paragraph with replacements and tab separators
    for i, new_text in enumerate(new_parts):
        if i > 0:
            paragraph.add_run("\t")  # Add tab back
        new_run = paragraph.add_run(new_text)

        # Apply formatting from the original part (if available)
        if i < len(paragraph.runs):
            original_run = paragraph.runs[i]
            new_run.bold = original_run.bold
            new_run.italic = original_run.italic
            new_run.underline = original_run.underline
            new_run.font.name = original_run.font.name
            if original_run.font.size:
                new_run.font.size = original_run.font.size
            if original_run.font.color:
                new_run.font.color.rgb = original_run.font.color.rgb


def populate_resume(
    person_id, db_path, template_path="template.docx", output_file="Resume.docx", filters=None
):
    """Loads a Word template and replaces placeholders with actual data."""
    # from docx.shared import Pt, Inches
    # from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document(template_path)
    data = fetch_resume_data(db_path, person_id, filters)

    # Replace placeholders while keeping formatting

    # Handle education, publications and employment sections
    for para in doc.paragraphs:
        if not para.runs:
            continue
        i = 0
        # original_style = para.style
        original_font = para.runs[0].font if para.runs else None
        if "{personal_info}" in para.text:
            para.clear()
            run = para.add_run(f"{data['full_name']}\n")
            run.bold = True
            contact = para.add_run(
                f"{data['email']}\t{data['linkedin']}\t{data['github']}\t"
            )
            if original_font:
                run.font.name = original_font.name
                run.font.size = original_font.size
                contact.font.name = original_font.name
                contact.font.size = Pt(original_font.size.pt - 6)

        if "{education}" in para.text:
            # Clear the paragraph while keeping its style
            para.clear()

            for degree, institution, graduation_year, gpa in data["education"]:
                run = para.add_run(f"{institution}, \t{graduation_year}\n")
                run.bold = True
                degree_run = para.add_run(f"\t{degree}\t GPA: {gpa}/4.0\n")

                # Preserve original font settings
                if original_font:
                    run.font.name = original_font.name
                    run.font.size = original_font.size
                    degree_run.font.name = original_font.name
                    degree_run.font.size = original_font.size

        elif "{publications}" in para.text:
            # Clear the paragraph while keeping its style
            para.clear()

            for title, authors, publication_date, venue, edition, pages in data[
                "publications"
            ]:
                run = para.add_run(
                    f"{authors}. ({publication_date}). {title}\n{venue}, {edition}, {pages}\n\n"
                )
                if original_font:
                    run.font.name = original_font.name
                    run.font.size = original_font.size

        elif "{employment}" in para.text:
            # Clear the paragraph
            current_para = para
            current_para.clear()
            prev_company = ""

            # Store the original paragraph format for resetting
            original_indent = para.paragraph_format.left_indent
            for position in data["employment"]:
                i += 1
                (
                    company,
                    location,
                    title,
                    start_date,
                    end_date,
                    field,
                    responsibilities,
                    # fields,
                ) = position

                # if not fields:
                #     print("No fields.")

                # Reset paragraph indentation for each new company
                if prev_company != company:
                    # Reset to original indent (usually 0 for left margin)
                    current_para.paragraph_format.left_indent = original_indent

                    # Add company with bold formatting
                    company_run = current_para.add_run(f"{company}")
                    company_run.bold = True
                    # Add location if available
                    location_run = current_para.add_run(f", {location}\n")

                    # Apply font formatting
                    if original_font:
                        company_run.font.size = Pt(original_font.size.pt + 2)
                        company_run.font.name = original_font.name
                        location_run.font.name = original_font.name
                        location_run.font.size  = original_font.size

                # Add job title and dates with tab spacing
                date_range = (
                    f"{start_date} - {end_date}" if end_date
                    else f"{start_date} - Present"
                )

                title_run = current_para.add_run(f"{title}\t{date_range}\n")

                if original_font:
                    title_run.font.name = original_font.name

                    title_run.font.size = original_font.size

                # Process responsibilities as bullet points

                if responsibilities:

                    for responsibility in responsibilities.split(";"):

                        if responsibility.strip():

                            # Create a properly formatted bullet point with the dash and spaces

                            bullet_run = current_para.add_run("•\t")

                            resp_run = current_para.add_run(

                                f"{responsibility.strip()}\n"

                            )

                            # Preserve font formatting

                            if original_font:
                                bullet_run.font.name = original_font.name

                                bullet_run.font.size = original_font.size

                                resp_run.font.name = original_font.name

                                resp_run.font.size = original_font.size


                # Update previous company

                prev_company = company
                if i < len(data["employment"]):
                    # Add space between job entries
                    current_para.add_run("\n")

        elif "{projects}" in para.text or "{personal_projects}" in para.text:
            if "{projects}" in para.text:
                focus = "projects"
            else:
                focus = "personal_projects"
            # Clear the paragraph

            current_para = para

            current_para.clear()

            # Store the original paragraph format for resetting

            original_indent = para.paragraph_format.left_indent
            for project in data[focus]:
                i += 1
                (
                    project_name, year, technologies, project_link, field, project_type, details

                ) = project
                print(i)
                print(len(data[focus]))
                print(field)

                print(project_link)

                print(project_type)

                # Reset paragraph indentation for each new company

                current_para.paragraph_format.left_indent = original_indent

                # Add company with bold formatting

                project_run = current_para.add_run(f"{project_name}")

                project_run.bold = True

                # Add location if available

                if field:
                    current_para.add_run(f", {field}, {year}\n")

                # Apply font formatting to company name

                if original_font:
                    project_run.font.name = original_font.name

                    project_run.font.size = original_font.size

                tools_run = current_para.add_run(f"Tools & Technologies: {technologies}\n")

                if original_font:
                    tools_run.font.name = original_font.name

                    tools_run.font.size = original_font.size

                # Process details as bullet points

                if details:

                    for detail in details.split(";"):

                        if detail.strip():

                            # Create a properly formatted bullet point with the dash and spaces

                            bullet_run = current_para.add_run("•\t")

                            detail_run = current_para.add_run(

                                f"{detail.strip()}\n"

                            )

                            # Preserve font formatting

                            if original_font:
                                bullet_run.font.name = original_font.name

                                bullet_run.font.size = original_font.size

                                detail_run.font.name = original_font.name

                                detail_run.font.size = original_font.size
                if i < len(data[focus]):
                    print(i)
                    # Add space between job entries
                    current_para.add_run("\n")

        elif "{professional_development}" in para.text:
            # Clear the paragraph
            current_para = para
            current_para.clear()

            # Store the original paragraph format for resetting
            original_indent = para.paragraph_format.left_indent

            for dev in data["professional_development"]:
                i += 1
                (
                    certification_name,
                    issuing_organization,
                    date_completed,
                    context,
                    field,
                    covered,
                ) = dev
                print(field)

                # Reset paragraph indentation for each new company

                current_para.paragraph_format.left_indent = original_indent

                # Add company with bold formatting
                prof_run = current_para.add_run(f"{certification_name}")
                prof_run.bold = True

                issue_run = current_para.add_run(f"-{issuing_organization} ({context} {date_completed}\n")

                # Apply font formatting to company name
                if original_font:
                    prof_run.font.name = original_font.name
                    prof_run.font.size = original_font.size
                    issue_run.font.name = original_font.name
                    issue_run.font.size = original_font.size

                # Process details as bullet points
                if covered:
                    for detail in covered.split(";"):
                        if detail.strip():
                            # Create a properly formatted bullet point with the dash and spaces
                            bullet_run = current_para.add_run("•\t")

                            detail_run = current_para.add_run(
                                f"{detail.strip()}\n"
                            )

                            # Preserve font formatting
                            if original_font:
                                bullet_run.font.name = original_font.name
                                bullet_run.font.size = original_font.size
                                detail_run.font.name = original_font.name
                                detail_run.font.size = original_font.size

                if i < len(data["professional_development"]):
                    # Add space between job entries
                    current_para.add_run("\n")

        elif "{technical_skills}" in para.text:
            i += 1
            # Clear the paragraph
            current_para = para
            current_para.clear()

            # Store the original paragraph format for resetting
            original_indent = para.paragraph_format.left_indent

            for skills in data["skills"]:
                (
                    skill,
                    details
                ) = skills

                # Reset paragraph indentation for each new company

                current_para.paragraph_format.left_indent = original_indent

                # Add company with bold formatting
                skill_run = current_para.add_run(f"{skill}: ")
                skill_run.bold = True

                # Apply font formatting to company name
                if original_font:
                    skill_run.font.name = original_font.name
                    skill_run.font.size = original_font.size


                # Process details as bullet points
                if details:
                    j = 0
                    for detail in details.split(";"):
                        if detail.strip():
                            # Create a properly formatted bullet point with the dash and spaces

                            detail_run = current_para.add_run(
                                f"{detail.strip()}"
                            )

                            # Preserve font formatting
                            if original_font:
                                detail_run.font.name = original_font.name
                                detail_run.font.size = original_font.size
                        if j < len(details) - 1:
                            current_para.add_run(',')

                if i < len(data["skills"]) - 1:
                    # Add space between job entries
                    current_para.add_run("\n")

    doc.save(output_file)
    print(f"Resume saved successfully as {output_file}")

filters = {"employment_filters": {
        "fields": ["Engineering"],
        "resp_fields": ["Problem Solving", "Technical Consultation", "Data Analysis","Automation ", "Optimization"],
    }}
populate_resume(1, db_path="resume.db", filters=filters)
